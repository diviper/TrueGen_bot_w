import os
from datetime import datetime
from pathlib import Path
from typing import Optional, Union, Dict, Any, List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell

from .models import ActData, ActItem
from .config import OUTPUT_DIR, TEMPLATES_DIR
from .logger import logger


class DocxGenerator:
    """Класс для генерации документов Word"""
    
    def __init__(self, template_path: Optional[Path] = None):
        """
        Инициализация генератора документов
        
        Args:
            template_path: Путь к шаблону документа (опционально)
        """
        self.template_path = template_path
    
    def generate_act(self, act_data: ActData, output_path: Optional[Path] = None) -> Path:
        """
        Генерирует документ акта
        
        Args:
            act_data: Данные акта
            output_path: Путь для сохранения файла (опционально)
            
        Returns:
            Path: Путь к сгенерированному файлу
            
        Raises:
            ValueError: Если не удалось сгенерировать документ
        """
        try:
            # Пытаемся использовать шаблон, если он указан и существует
            if self.template_path and self.template_path.exists():
                try:
                    doc = Document(self.template_path)
                    logger.info(f'Использован шаблон: {self.template_path}')
                except Exception as e:
                    logger.warning(f'Ошибка при загрузке шаблона {self.template_path}: {e}. Используется стандартный шаблон.')
                    doc = self._create_default_document()
            else:
                doc = self._create_default_document()
            
            # Заменяем плейсхолдеры в документе (если есть)
            self._replace_placeholders(doc, act_data)
            
            # Добавляем заголовок, если его нет в шаблоне
            if not self._has_title(doc):
                self._add_header(doc, act_data)
            
            # Добавляем таблицу с позициями
            self._add_items_table(doc, act_data.items)
            
            # Добавляем итоговую сумму, если её нет
            if not self._has_total(doc):
                self._add_total(doc, act_data.total)
            
            # Сохраняем документ
            return self._save_document(doc, act_data, output_path)
            
        except Exception as e:
            logger.error(f'Ошибка при генерации акта: {e}', exc_info=True)
            raise ValueError(f'Не удалось сгенерировать документ: {e}')
    
    def _create_default_document(self) -> Document:
        """Создает новый документ с настройками по умолчанию"""
        doc = Document()
        self._apply_default_styles(doc)
        return doc
        
    def _apply_default_styles(self, doc: Document) -> None:
        """Применяет стили по умолчанию к документу"""
        # Устанавливаем шрифт по умолчанию
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Настраиваем стили заголовков
        for level in range(1, 6):
            try:
                heading_style = doc.styles[f'Heading {level}']
                heading_font = heading_style.font
                heading_font.name = 'Times New Roman'
                heading_font.bold = True
                heading_font.size = Pt(16 - (level * 2))
            except KeyError:
                continue
    
    def _has_title(self, doc: Document) -> bool:
        """Проверяет, есть ли в документе заголовок"""
        if not doc.paragraphs:
            return False
            
        # Ищем заголовок в первых 3 параграфах
        for para in doc.paragraphs[:3]:
            if 'акт' in para.text.lower() and 'выполнен' in para.text.lower():
                return True
        return False
    
    def _has_total(self, doc: Document) -> bool:
        """Проверяет, есть ли в документе строка с итоговой суммой"""
        if not doc.paragraphs:
            return False
            
        # Ищем строку с итогом в последних 5 параграфах
        for para in doc.paragraphs[-5:]:
            if 'итог' in para.text.lower():
                return True
        return False
    
    def _replace_placeholders(self, doc: Document, act_data: ActData) -> None:
        """Заменяет плейсхолдеры в документе на реальные данные"""
        placeholders = {
            '{date}': act_data.date.strftime('%d.%m.%Y'),
            '{object}': act_data.object_name,
            '{total}': f'{act_data.total:.2f} ₽'
        }
        
        for para in doc.paragraphs:
            for old_text, new_text in placeholders.items():
                if old_text in para.text:
                    self._replace_text_in_paragraph(para, old_text, new_text)
        
        # Заменяем в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in placeholders.items():
                            if old_text in para.text:
                                self._replace_text_in_paragraph(para, old_text, new_text)
    
    def _replace_text_in_paragraph(self, para, old_text: str, new_text: str) -> None:
        """Заменяет текст в параграфе с сохранением форматирования"""
        if old_text not in para.text:
            return
            
        # Сохраняем стиль первого run'а
        inline = para.runs
        for i in range(len(inline)):
            if old_text in inline[i].text:
                text = inline[i].text.replace(old_text, new_text)
                inline[i].text = text
                break
    
    def _add_header(self, doc: Document, act_data: ActData) -> None:
        """Добавляет заголовок акта"""
        # Добавляем заголовок, если его еще нет
        if not self._has_title(doc):
            title = doc.add_heading('АКТ ВЫПОЛНЕННЫХ РАБОТ', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем дату и объект, если их еще нет
        date_found = any('дата:' in p.text.lower() for p in doc.paragraphs[:5])
        if not date_found:
            doc.add_paragraph(f"Дата: {act_data.date.strftime('%d.%m.%Y')}")
        
        object_found = any('объект:' in p.text.lower() for p in doc.paragraphs[:5])
        if not object_found:
            doc.add_paragraph(f"Объект: {act_data.object_name}")
        
        # Добавляем пустую строку, если нужно
        if doc.paragraphs and doc.paragraphs[-1].text.strip() != '':
            doc.add_paragraph()
    
    def _add_items_table(self, doc: Document, items: List[ActItem]) -> None:
        """Добавляет таблицу с позициями"""
        # Ищем существующую таблицу с позициями
        table_to_fill = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if any(keyword in cell.text.lower() for keyword in ['наименование', 'кол-во', 'ед.', 'цена', 'сумма']):
                        table_to_fill = table
                        break
                if table_to_fill:
                    break
            if table_to_fill:
                break
        
        if table_to_fill:
            # Заполняем существующую таблицу
            self._fill_existing_table(table_to_fill, items)
        else:
            # Создаем новую таблицу
            self._create_new_table(doc, items)
    
    def _fill_existing_table(self, table: Table, items: List[ActItem]) -> None:
        """Заполняет существующую таблицу данными"""
        # Находим индексы колонок
        headers = {}
        for i, cell in enumerate(table.rows[0].cells):
            text = cell.text.strip().lower()
            if 'наименование' in text:
                headers['name'] = i
            elif 'кол-во' in text or 'количество' in text:
                headers['qty'] = i
            elif 'ед.' in text or 'единица' in text:
                headers['unit'] = i
            elif 'цена' in text:
                headers['price'] = i
            elif 'сумма' in text or 'стоимость' in text:
                headers['total'] = i
        
        # Добавляем строки с данными
        for idx, item in enumerate(items, 1):
            row_cells = table.add_row().cells
            for col, cell in enumerate(row_cells):
                if col == headers.get('name', 0):
                    cell.text = item.name
                elif col == headers.get('qty', 1):
                    cell.text = str(item.quantity)
                elif col == headers.get('unit', 2):
                    cell.text = item.unit
                elif col == headers.get('price', 3):
                    cell.text = f"{item.price:.2f}"
                elif col == headers.get('total', 4):
                    cell.text = f"{item.total:.2f}"
    
    def _create_new_table(self, doc: Document, items: List[ActItem]) -> None:
        """Создает новую таблицу с данными"""
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Заголовки столбцов
        hdr_cells = table.rows[0].cells
        headers = ['№', 'Наименование', 'Кол-во', 'Ед.', 'Сумма, ₽']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Добавляем строки с позициями
        for idx, item in enumerate(items, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = item.name
            row_cells[2].text = str(item.quantity)
            row_cells[3].text = item.unit
            row_cells[4].text = f"{item.total:.2f}"
        
        # Добавляем отступ после таблицы
        doc.add_paragraph()
    
    def _add_total(self, doc: Document, total: float) -> None:
        """Добавляет итоговую сумму"""
        # Проверяем, есть ли уже строка с итогом
        if self._has_total(doc):
            return
            
        # Добавляем отступ, если нужно
        if doc.paragraphs and doc.paragraphs[-1].text.strip() != '':
            doc.add_paragraph()
            
        # Добавляем итоговую сумму
        p = doc.add_paragraph()
        p.add_run(f"Итого: {total:.2f} ₽").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Добавляем примечание о НДС
        p = doc.add_paragraph()
        p.add_run("Без НДС").italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _save_document(self, doc: Document, act_data: ActData, output_path: Optional[Path] = None) -> Path:
        """
        Сохраняет документ в файл
        
        Args:
            doc: Документ для сохранения
            act_data: Данные акта
            output_path: Путь для сохранения (опционально)
            
        Returns:
            Path: Путь к сохраненному файлу
            
        Raises:
            IOError: Если не удалось сохранить файл
        """
        try:
            # Создаем директорию output, если её нет
            if output_path is None:
                output_path = self._get_default_output_path(act_data)
            
            output_dir = output_path.parent
            try:
                output_dir.mkdir(parents=True, exist_ok=True)
                logger.info(f'Создана директория для сохранения: {output_dir.absolute()}')
            except Exception as e:
                error_msg = f'Не удалось создать директорию {output_dir}: {e}'
                logger.error(error_msg)
                raise IOError(error_msg)
            
            # Сохраняем документ
            doc.save(str(output_path))
            logger.info(f'Документ успешно сохранен: {output_path.absolute()}')
            return output_path
            
        except Exception as e:
            error_msg = f'Ошибка при сохранении документа: {e}'
            logger.error(error_msg, exc_info=True)
            raise IOError(error_msg)


def generate_act_document(act_data: ActData, template_path: Optional[Path] = None, output_path: Optional[Path] = None) -> Path:
    """
    Генерирует документ акта
    
    Args:
        act_data: Данные акта
        template_path: Путь к шаблону (опционально)
        output_path: Путь для сохранения (опционально)
        
    Returns:
        Path: Путь к сгенерированному файлу
        
    Raises:
        Exception: Если произошла ошибка при генерации документа
    """
    try:
        logger.info(f'Создание генератора документов. Шаблон: {template_path}')
        generator = DocxGenerator(template_path)
        logger.info('Генератор документов создан успешно')
        
        logger.info('Начало генерации акта...')
        result_path = generator.generate_act(act_data, output_path)
        logger.info(f'Акт успешно сгенерирован: {result_path}')
        
        return result_path
    except Exception as e:
        logger.error(f'Ошибка при генерации документа: {str(e)}', exc_info=True)
        raise