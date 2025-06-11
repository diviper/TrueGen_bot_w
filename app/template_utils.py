import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional

from .logger import logger

def create_default_template(output_path: Optional[Path] = None) -> Path:
    """
    Создает шаблон документа по умолчанию
    
    Args:
        output_path: Путь для сохранения шаблона. Если не указан, используется 'templates/act_template.docx'
        
    Returns:
        Path: Путь к созданному шаблону
    """
    if output_path is None:
        output_path = Path('templates/act_template.docx')
    
    # Создаем директорию, если её нет
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Заголовок
    title = doc.add_heading('АКТ ВЫПОЛНЕННЫХ РАБОТ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата и объект (заполняются автоматически)
    doc.add_paragraph('Дата: {date}')
    doc.add_paragraph('Объект: {object_name}')
    doc.add_paragraph()  # Пустая строка
    
    # Таблица с позициями
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Заголовки столбцов
    hdr_cells = table.rows[0].cells
    headers = ['№', 'Наименование', 'Кол-во', 'Ед.', 'Сумма, ₽']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Сохраняем шаблон
    doc.save(output_path)
    logger.info(f'Создан новый шаблон: {output_path.absolute()}')
    return output_path

def is_valid_docx(file_path: Path) -> bool:
    """
    Проверяет, является ли файл корректным DOCX
    
    Args:
        file_path: Путь к файлу для проверки
        
    Returns:
        bool: True если файл корректен, иначе False
    """
    try:
        Document(file_path)
        return True
    except Exception as e:
        logger.error(f'Ошибка при проверке файла {file_path}: {str(e)}')
        return False

def ensure_template_exists(template_path: Optional[Path] = None) -> Path:
    """
    Проверяет и создает шаблон, если его нет или он поврежден
    
    Args:
        template_path: Путь к шаблону. Если не указан, используется 'templates/act_template.docx'
        
    Returns:
        Path: Путь к существующему шаблону
    """
    if template_path is None:
        template_path = Path('templates/act_template.docx')
    
    if not template_path.exists() or not is_valid_docx(template_path):
        return create_default_template(template_path)
    return template_path
